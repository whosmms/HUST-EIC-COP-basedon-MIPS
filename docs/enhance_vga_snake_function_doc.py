import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOC = ROOT / "docs" / "vga_snake_function_analysis.docx"
OUT_DOC = ROOT / "docs" / "vga_snake_function_analysis_commented.docx"


def set_run_font(run, name="Microsoft YaHei", size=9.0, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_para_spacing(paragraph, before=0, after=4, line=1.05):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_h1(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=10, after=5, line=1.0)
    r = p.add_run(text)
    set_run_font(r, size=15, bold=True, color=(31, 78, 121))


def add_h2(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=7, after=4, line=1.0)
    r = p.add_run(text)
    set_run_font(r, size=12.5, bold=True, color=(47, 84, 150))


def add_body(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p)
    r = p.add_run(text)
    set_run_font(r, size=10.5)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        p = table.rows[0].cells[i].paragraphs[0]
        r = p.add_run(header)
        set_run_font(r, size=9.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(value))
            set_run_font(r, size=9.0)
    doc.add_paragraph()


def comment_prefix(path_text):
    return "# 注释：" if path_text and path_text.endswith(".xdc") else "// 注释："


def extract_identifier(line):
    s = line.strip().rstrip(",;")
    if "=" in s:
        left = s.split("=", 1)[0].strip()
    else:
        left = s
    parts = left.split()
    return parts[-1] if parts else ""


def explain_parameter(name):
    meanings = {
        "COLS": "把游戏棋盘宽度定义成 40 格；后续像素坐标会除以 16 映射到这些列上。",
        "ROWS": "把游戏棋盘高度定义成 30 格；40x30 正好对应 640x480 中每格 16 像素。",
        "MAXLEN": "限制玩家蛇最多占用多少格，防止数组越界和资源无限增长。",
        "MAX_ENEMIES": "限制同时存在的敌人数，避免敌人数组超过综合时固定的硬件规模。",
        "ENEMY_MIN_LEN": "限定敌人最短长度，保证生成出来的敌人不是单个点，能形成障碍。",
        "ENEMY_MAX_LEN": "限定敌人最长长度，保证敌人身体数组大小固定、可综合。",
        "ENEMY_SPAWN_MIN_TICKS": "给敌人生成设置最小等待时间，防止敌人连续刷出导致游戏不可玩。",
        "ENEMY_SPAWN_RANDOM_MASK": "只取 seed 的低若干位作为额外等待时间，使生成间隔随机但仍有上限。",
        "FOOD_SPAWN_MIN_TICKS": "给食物生成设置最小等待时间，避免食物刚吃完就立刻满屏刷新。",
        "FOOD_SPAWN_RANDOM_MASK": "限制食物随机延迟的最大扰动范围，让刷新有随机性但节奏可控。",
        "INITIAL_LIVES": "设置开局生命数；碰撞后从这里开始扣减。",
        "MAX_LIVES": "限制蓝色生命食物最多能把生命加到多少，避免生命无限增长。",
        "INVINCIBLE_TICKS": "规定受伤后的无敌持续 tick 数，用来给玩家恢复和重生缓冲。",
        "CELL_W": "每个格子横向 16 像素，因此 pixel_x 的高位可以直接作为列号。",
        "CELL_H": "每个格子纵向 16 像素，因此 pixel_y 的高位可以直接作为行号。",
        "TICK_DIV": "把 25MHz 像素时钟分频成蛇移动节拍；游戏移动速度由它控制。",
        "BUTTON_DEBOUNCE_CLKS": "要求按键电平稳定一段时间才确认，过滤机械按键抖动。",
        "ONE_SECOND_CLKS": "把系统时钟分频成 1 秒，用来统计游戏时间。",
        "SEVENSEG_REFRESH_DIV": "控制数码管扫描刷新频率，让 8 位数码管看起来同时点亮。",
    }
    return meanings.get(name)


def explain_localparam(name, line):
    meanings = {
        "EMPTY": "把空格编码成 000，后面渲染模块看到这个编码就画背景和网格线。",
        "BODY": "把玩家身体编码成 001，渲染时显示为绿色身体。",
        "HEAD": "把玩家头编码成 010，渲染时显示为白色，方便和身体区分。",
        "RED_FOOD": "把红色食物编码成 011，吃到后加 1 分并增长 1 节。",
        "YELLOW_FOOD": "把黄色食物编码成 100，吃到后加 2 分并触发额外增长。",
        "BLUE_FOOD": "把生命食物编码成 101，渲染为蓝色，吃到后加生命。",
        "ENEMY_HEAD": "把敌人头编码成 110，渲染为亮品红，碰到它会伤害玩家。",
        "ENEMY_BODY": "把敌人身体编码成 111，渲染为暗品红，撞到身体会击杀敌人并加分。",
        "FOOD_RED": "内部用 0 表示红色食物类型，和 cell_state 的显示编码分开管理。",
        "FOOD_YELLOW": "内部用 1 表示黄色食物类型，便于 select_food_type 返回。",
        "FOOD_GREEN": "内部用 2 表示生命食物类型；变量名沿用 GREEN，但显示时是蓝色生命食物。",
        "DIR_UP": "用 00 表示向上，方向只需 2 位即可覆盖上下左右。",
        "DIR_DOWN": "用 01 表示向下，与 DIR_UP 组成相反方向判断。",
        "DIR_LEFT": "用 10 表示向左。",
        "DIR_RIGHT": "用 11 表示向右；开局默认向右移动。",
        "GRID_CELLS": "棋盘总格数，后续随机位置用 lfsr % GRID_CELLS 映射到合法格号。",
        "INIT_HEAD_IDX": "把初始蛇头放在棋盘中心，保证开局离墙较远。",
        "INIT_BODY1_IDX": "初始身体第一节放在蛇头左边，使蛇开局朝右时身体跟在后面。",
        "INIT_BODY2_IDX": "初始身体第二节继续向左，形成长度为 3 的初始蛇。",
        "INVALID_INDEX": "用 16'hFFFF 表示无效格子，因为正常格号小于 40x30。",
        "RESP_SCAN": "玩家重生流程的扫描阶段：先寻找可放置蛇头的位置。",
        "RESP_BUILD": "玩家重生流程的构造阶段：找到头后继续摆放身体。",
        "NO_ENEMY_SLOT": "用 MAX_ENEMIES 作为无空敌人槽标记，因为合法槽号只到 MAX_ENEMIES-1。",
    }
    if name in meanings:
        return meanings[name]
    if name.startswith("H_") or name.startswith("V_"):
        return "这是 VGA 640x480 时序参数，用可见区、前肩、同步脉冲、后肩共同满足显示器扫描规范。"
    if name.startswith("GAMEOVER_TEXT_"):
        return "规定 GAMEOVER 点阵文字在 640x480 画面中的位置或尺寸，使文字居中显示。"
    return None


def explain_declaration(name):
    meanings = {
        "snake_map": "用 1200 位位图快速判断某格是否被玩家身体占用，碰撞和生成检查会更直接。",
        "snake_pos": "保存蛇每一节的格子编号；移动时通过更新头尾指针形成身体跟随效果。",
        "slen": "记录当前蛇长，吃食物时增加，越界时判定异常。",
        "head_ptr": "指向 snake_pos 中当前蛇头位置；使用环形数组可以避免整体搬移玩家身体。",
        "tail_ptr": "指向 snake_pos 中当前蛇尾位置；普通移动时尾巴要释放这个格子。",
        "head_idx": "保存蛇头所在格号，碰撞、渲染和食物判断都以它为核心。",
        "food_idx": "保存红色食物所在格号，玩家新头等于它时触发加分增长。",
        "food_valid": "标记红色食物是否存在，避免同一种食物重复生成。",
        "yellow_food_idx": "保存黄色食物所在格号。",
        "yellow_food_valid": "标记黄色食物是否存在。",
        "green_food_idx": "保存生命食物所在格号；虽然变量叫 green，显示中对应蓝色生命食物。",
        "green_food_valid": "标记生命食物是否存在。",
        "pending_food": "表示当前正在寻找食物生成位置；扫描到空格后才真正放置食物。",
        "pending_food_type": "保存本次准备生成的食物类型，扫描期间类型不再变化。",
        "food_scan_idx": "食物生成从随机格开始逐格扫描，直到找到不冲突的位置。",
        "food_spawn_countdown": "食物生成倒计时；归零才提出生成请求，让食物出现有节奏。",
        "invincible_ticks": "受伤后的无敌剩余时间，非零时部分碰撞不会再次扣命。",
        "invincible_blink": "控制无敌期间玩家闪烁显示，给用户一个受伤反馈。",
        "growth_pending": "黄色食物带来的延迟增长次数；后续移动时不删尾巴完成额外变长。",
        "hx": "蛇头的 x 坐标缓存，避免每次移动都从格号除法恢复。",
        "hy": "蛇头的 y 坐标缓存，避免每次移动都从格号除法恢复。",
        "cur_dir": "记录当前实际移动方向，用来禁止玩家直接反向掉头。",
        "next_dir": "记录已经接收但等待下一 tick 生效的方向。",
        "lfsr": "16 位伪随机寄存器，是 seed 的主要来源。",
        "enemy_alive": "记录每个敌人槽是否正在使用；生成时找空槽，击杀时清零。",
        "enemy_dir": "保存每个敌人当前方向，避免随机移动时直接 180 度回头。",
        "enemy_len": "保存每个敌人的实际长度，使不同敌人可以有随机长度。",
        "enemy_pos": "保存所有敌人的头和身体格号，布局为 槽号*最大长度+段号。",
        "enemy_spawn_countdown": "敌人生成倒计时，避免敌人一开始或击杀后立刻刷满。",
        "pending_enemy_spawn": "表示正在寻找敌人生成位置，找到合法连续空间后才生成。",
        "enemy_scan_idx": "敌人生成的扫描格号；从随机起点开始，找可容纳整条敌人的连续空间。",
        "rand_dir": "从 lfsr 取出的候选敌人方向。",
        "try_dir": "经过反向修正后的实际尝试方向。",
        "enemy_next_idx": "敌人头如果按 try_dir 移动，将到达的下一格。",
        "player_gameover": "本 tick 内的临时标志，用来阻止游戏结束后继续处理敌人和食物。",
        "player_grew": "记录玩家本 tick 是否增长，敌人撞玩家身体时要考虑尾巴是否会离开。",
        "player_hit_enemy_body": "记录玩家是否撞到敌人身体，撞身体会击杀敌人而不是伤害玩家。",
        "killed_enemy_slot": "保存被玩家撞身体击杀的敌人槽号。",
        "player_ate_food": "记录玩家新头是否落在食物格上。",
        "eaten_food_type": "记录吃到的是哪种食物，以决定加分、增长或加生命。",
        "player_damaged": "记录玩家本 tick 是否受伤，受伤时暂停敌人继续结算。",
        "spawn_slot": "保存本次准备生成敌人的空槽号。",
        "spawn_len": "保存本次敌人的随机长度。",
        "enemy_hit_player": "记录敌人本次移动是否撞到玩家头或身体。",
        "enemy_can_move": "表示敌人尝试方向是否可走，不可走时本 tick 不更新该敌人。",
        "cell_index": "把渲染查询的 cell_x/cell_y 转成一维格号，和蛇/食物/敌人数组统一比较。",
        "show_player": "无敌闪烁时控制玩家是否画出来，实现受伤反馈。",
        "tick_cnt": "累积像素时钟，用来生成较慢的游戏移动 tick。",
        "tick_q": "移动 tick 的单周期脉冲，蛇和敌人只在它为 1 时前进一步。",
        "elapsed_min": "记录游戏运行分钟数，用于数码管显示和提高敌人上限。",
        "elapsed_sec": "记录游戏运行秒数，用于数码管显示。",
        "active_enemy_limit": "根据时间决定允许存在的敌人数，2 分钟后提高难度。",
    }
    return meanings.get(name)


def explain_line(path_text, title, no, code):
    s = code.strip()
    if not s:
        return None
    if "注释：" in s:
        return None

    # Very specific lines where the reason matters most.
    specific = {
        "bucket = seed % 10;": "把 16 位 seed 压成 10 个桶，是为了用整数比较实现概率：10 个桶里 5 个红、3 个黄、2 个生命食物。",
        "if (bucket < 5)": "桶号 0-4 占 5/10，所以红色食物概率最高，作为普通奖励最常见。",
        "else if (bucket < 8)": "桶号 5-7 占 3/10，所以黄色食物比红色少见，但仍比较容易出现。",
        "else": "剩下桶号 8-9 占 2/10，所以生命食物最少见，避免生命补给过多降低难度。",
        "random_enemy_len = ENEMY_MIN_LEN + (seed % span);": "先用取模把 seed 映射到 0..span-1，再加最小长度；这样敌人长度一定落在允许范围内，又保留随机差异。",
        "next_spawn_delay = ENEMY_SPAWN_MIN_TICKS + (seed & ENEMY_SPAWN_RANDOM_MASK);": "最小等待时间保证不会连续刷敌人，低位掩码提供额外随机等待，让敌人出现节奏不可预测但有上限。",
        "next_food_delay = FOOD_SPAWN_MIN_TICKS + (seed & FOOD_SPAWN_RANDOM_MASK);": "食物也采用“固定最小间隔+随机扰动”，这样既不会太密集，也不会每次间隔完全一样。",
        "lfsr <= 16'hACE1;": "给伪随机发生器一个非零初值；LFSR 不能从全 0 启动，否则反馈后仍可能一直为 0。",
        "lfsr <= {lfsr[14:0], lfsr[15] ^ lfsr[13] ^ lfsr[12] ^ lfsr[10]};": "把旧随机数左移，并用多个抽头异或产生新位；这样硬件只需移位和异或就能持续生成伪随机序列。",
        "enemy_scan_idx <= lfsr % GRID_CELLS;": "敌人生成不是从固定 0 号格扫描，而是从随机格开始扫描，避免敌人总在同一区域出现。",
        "food_scan_idx <= lfsr % GRID_CELLS;": "食物生成从随机格开始找空位，避免食物位置总偏向左上角。",
        "pending_food_type <= select_food_type(lfsr ^ head_idx);": "把蛇头位置混入 lfsr，表示玩家当前局面会扰动下一次食物类型，减少固定序列感。",
        "spawn_len = random_enemy_len(lfsr ^ enemy_scan_idx);": "把敌人生成起点混入 lfsr，同一时刻不同扫描点也能得到不同长度。",
        "enemy_spawn_countdown <= next_spawn_delay(lfsr ^ enemy_scan_idx);": "敌人成功生成后，用生成位置扰动下一次等待时间，避免生成周期固定。",
        "food_spawn_countdown <= next_food_delay(lfsr ^ food_scan_idx);": "食物生成成功后，用本次食物位置扰动下一次等待时间，让位置和节奏都变化。",
        "pending_enemy_spawn <= 1'b1;": "倒计时归零后先进入“待生成”状态，因为还必须找到不撞玩家、食物和敌人的连续空格。",
        "pending_food <= 1'b1;": "倒计时归零后先进入“待生成”状态，因为要扫描一个没有被蛇、敌人、已有食物占用的格子。",
        "new_head_ptr = prev_snake_index(head_ptr);": "环形数组中把头指针向前挪一格，用新槽保存下一步蛇头；这样不用搬移整条蛇。",
        "next_tail_ptr = prev_snake_index(tail_ptr);": "普通移动时尾巴也要前移一节，表现为蛇整体向前而长度不变。",
        "snake_map[tailidx] <= 1'b0;": "释放旧尾巴所在格，否则下一次碰撞检测会误以为尾巴还占在那里。",
        "snake_map[head_idx] <= 1'b1;": "旧头在移动后变成身体，所以要加入身体位图，供碰撞和渲染使用。",
        "head_idx <= nidx;": "把蛇头更新到新格，之后渲染和碰撞都以这个新头为准。",
        "tailidx = snake_pos[tail_ptr];": "先记住当前尾巴位置，因为普通移动会释放它，自撞判断也允许走到即将离开的尾巴格。",
        "if ((nidx == head_idx) || (snake_map[nidx] && nidx != tailidx)) begin": "新头撞到当前头或身体就是自撞；但如果撞到的是本 tick 会离开的尾巴，则允许通过。",
        "if (food_valid && (nidx == food_idx)) begin": "新头进入红食物格才算吃到红食物，旧身体经过食物不触发奖励。",
        "end else if (yellow_food_valid && (nidx == yellow_food_idx)) begin": "红食物没吃到时再判断黄色，保证同一格只触发一种食物效果。",
        "end else if (green_food_valid && (nidx == green_food_idx)) begin": "最后判断生命食物，避免和红/黄食物处理互相覆盖。",
        "if (growth_pending != 0) begin": "有延迟增长时本步不删尾巴，用“多保留一节”的方式实现继续变长。",
        "growth_pending <= growth_pending - 1'b1;": "每完成一次延迟增长就消耗一次计数，防止黄色食物造成无限增长。",
        "if (lives < MAX_LIVES)": "吃生命食物前检查上限，避免蓝色食物把生命堆到显示和设计范围之外。",
        "lives <= lives + 1'b1;": "生命食物不改变蛇长，而是增加容错次数。",
        "enemy_pos[enemy_base + ej] <= enemy_pos[enemy_base + ej - 1];": "敌人身体每一节复制前一节的位置，形成“身体跟随敌人头”的移动效果。",
        "enemy_pos[enemy_base] <= enemy_next_idx;": "身体跟随完成后，敌人头再进入新格，整条敌人才完成一次前进。",
        "if (is_opposite_dir(enemy_dir[ei], rand_dir))": "敌人也禁止直接反向，因为直接掉头会让头撞进自己的身体，移动效果不自然。",
        "try_dir = turn_right(enemy_dir[ei]);": "随机方向若是反向，就改成右转，既避开非法回头，又保持敌人有转向行为。",
        "enemy_can_move = enemy_move_valid(ei[2:0], try_dir);": "敌人移动前先确认目标格不越界、不撞敌人、不占食物；不可走时本 tick 保持原位。",
        "enemy_hit_player =": "敌人移动后的新头需要和玩家新头、旧头、身体同时比较，才能覆盖同向相撞和敌人撞身体等情况。",
        "if (enemy_hit_player && (invincible_ticks == 0)) begin": "只有玩家不在无敌期时，敌人撞玩家才扣命；无敌期避免连续掉命。",
        "assign cell_state =": "这里用优先级把一个格子翻译成显示状态；排在前面的对象会覆盖排在后面的对象。",
        "(!show_cells) ? EMPTY :": "空闲或越界时显示为空，避免未开始游戏时残留画面。",
        "(show_player && (cell_index == head_idx)) ? HEAD :": "玩家头优先级最高之一，保证蛇头不会被身体或食物颜色盖住。",
        "enemy_head_cell ? ENEMY_HEAD :": "敌人头单独显示，玩家看到它就知道这是危险碰撞点。",
        "enemy_body_cell ? ENEMY_BODY :": "敌人身体放在食物之后显示，表示食物优先显示但敌人身体仍参与碰撞。",
        "(show_player && snake_map[cell_index]) ? BODY :": "只有无敌闪烁允许显示玩家时才画身体，实现受伤闪烁效果。",
        "assign cell_x = pixel_x[9:4];": "每格 16 像素，所以取 pixel_x 高位等价于除以 16；硬件中这样比除法更省资源。",
        "assign cell_y = pixel_y[8:4];": "每格 16 像素，所以取 pixel_y 高位等价于除以 16，把像素行映射成棋盘行。",
        "wire border = (local_x == 0) || (local_y == 0);": "每个格子的第 0 行/列画灰线，玩家能看清 40x30 网格边界。",
        "if (!active) begin": "VGA 消隐区不能输出有效图像内容，所以 RGB 清零，防止显示器看到脏色。",
        "rgb_r <= 3'b000; rgb_g <= 3'b111; rgb_b <= 3'b000;": "绿色用于玩家身体，和黑色背景、品红敌人形成明显区分。",
        "rgb_r <= 3'b111; rgb_g <= 3'b111; rgb_b <= 3'b111;": "白色用于玩家头或 GAMEOVER 文字，亮度最高，最容易被看到。",
        "rgb_r <= 3'b111; rgb_g <= 3'b000; rgb_b <= 3'b000;": "红色表示普通食物，颜色醒目且符合奖励提示习惯。",
        "rgb_r <= 3'b111; rgb_g <= 3'b111; rgb_b <= 3'b000;": "黄色表示更高价值食物，和红色普通食物区分。",
        "rgb_r <= 3'b000; rgb_g <= 3'b000; rgb_b <= 3'b111;": "蓝色表示生命食物，和红黄奖励食物区分为补给类道具。",
        "rgb_r <= 3'b111; rgb_g <= 3'b000; rgb_b <= 3'b111;": "亮品红用于敌人头，提醒玩家这是危险目标。",
        "rgb_r <= 3'b011; rgb_g <= 3'b000; rgb_b <= 3'b011;": "暗品红用于敌人身体，既能看出敌人形状，也能和敌人头区分。",
        "assign active_video = (hcnt < H_VISIBLE) && (vcnt < V_VISIBLE);": "只有扫描点位于 640x480 可见区时才允许渲染，前肩/同步/后肩阶段属于消隐。",
        "hsync <= 0;": "VGA 同步脉冲为低有效，拉低告诉显示器一行扫描结束并准备回扫。",
        "vsync <= 0;": "VGA 场同步脉冲为低有效，拉低告诉显示器一帧扫描结束并准备回到顶部。",
        "pixel_x <= hcnt;": "可见区内把水平计数作为像素 x 坐标，渲染模块据此查询格子。",
        "pixel_y <= vcnt;": "可见区内把垂直计数作为像素 y 坐标，渲染模块据此查询格子。",
        "tick_q <= 1'b1;": "分频计数到达上限时只拉高一个周期，驱动蛇和敌人前进一步。",
        "tick_q <= 1'b0;": "非移动周期保持 0，防止蛇在 25MHz 像素时钟下飞速移动。",
        "wire start = (state == 2'd0) ? dir_req_valid : 1'b0;": "只有空闲状态下方向键才启动游戏，运行中方向键只改变移动方向。",
        "wire pause  = pause_reg;": "把中键切换得到的寄存器状态送入游戏逻辑，统一控制暂停。",
        "wire reset_game = reset_game_reg;": "GAMEOVER 时按中键产生单周期复位请求，让 snake_game 回到空闲状态。",
        "assign gpio_status = {lives, status_state, score};": "把生命、状态、分数打包给 AXI GPIO/MicroBlaze，SDK 程序可以读到硬件游戏状态。",
        "wire gameover_text_on =": "只有状态为 GAMEOVER 且像素落在字形点阵上时才点亮文字。",
        "assign vga_r = (state == 2'd3) ? (gameover_text_on ? 3'b111 : 3'b000) : rgb_r;": "游戏结束时红色通道由 GAMEOVER 文字接管，其他区域清黑。",
        "assign vga_g = (state == 2'd3) ? (gameover_text_on ? 3'b111 : 3'b000) : rgb_g;": "游戏结束时绿色通道同样只在文字点亮处输出，和红蓝一起形成白字。",
        "assign vga_b = (state == 2'd3) ? (gameover_text_on ? 3'b111 : 3'b000) : rgb_b;": "游戏结束时蓝色通道同样参与白字输出，保证 GAMEOVER 清楚显示。",
    }
    if s in specific:
        return specific[s]

    # Seed-related calls with expressions.
    if "input [15:0] seed" in s or "input [15:0]  seed" in s:
        return "seed 是函数的随机输入口；它不来自外部管脚，而是在调用处由当前 lfsr 或 lfsr 的扰动值传进来。"
    if "weighted_food_type = FOOD_RED" in s:
        return "桶号落在红色范围时返回红食物，代表本次随机选择普通加分增长奖励。"
    if "weighted_food_type = FOOD_YELLOW" in s:
        return "桶号落在黄色范围时返回黄食物，代表本次随机选择更高分且额外增长的奖励。"
    if "weighted_food_type = FOOD_GREEN" in s:
        return "桶号落在最后 2 个桶时返回生命食物，让补生命较稀有，游戏仍有挑战。"
    if "chosen = weighted_food_type(seed)" in s:
        return "先按概率得到候选食物类型，后面再检查这种食物当前是否已经存在。"
    if "if (food_type_available(chosen))" in s:
        return "如果候选类型当前没在场上，就直接采用它，保持随机权重结果。"
    if "select_food_type = FOOD_RED" in s:
        return "候选类型不可用时优先补一个当前不存在的红食物，保证场上食物种类尽量齐全。"
    if "select_food_type = FOOD_YELLOW" in s:
        return "红食物已存在而黄食物不存在时，改生成黄食物。"
    if "select_food_type = FOOD_GREEN" in s:
        return "红黄都存在时，只剩生命食物可生成，因此选择它。"

    # Parameters and constants.
    if s.startswith("parameter "):
        meaning = explain_parameter(extract_identifier(s))
        return meaning or "这个参数把可调数值集中放在模块开头，便于改变游戏规模或速度而不改主体逻辑。"
    if s.startswith("localparam "):
        meaning = explain_localparam(extract_identifier(s), s)
        return meaning or "内部常量把魔法数字命名，后面逻辑读起来能看出它代表的游戏含义。"

    # Declarations with concrete meanings.
    if s.startswith("reg ") or s.startswith("wire "):
        name = extract_identifier(s)
        # Array declarations like enemy_alive [0:...].
        name = name.split("[", 1)[0].strip()
        meaning = explain_declaration(name)
        if meaning:
            return meaning
        if "_idx" in name:
            return "带 idx 的变量通常保存一维格号，用来把二维棋盘统一成数组下标比较。"
        if "_valid" in name:
            return "valid 标志表示对应对象是否真的存在，避免用旧坐标误判碰撞或显示。"

    # Game logic patterns.
    if "spawn_cell_free =" in s:
        return "生成敌人前必须确认目标格可用，否则敌人可能刷在玩家、食物或其他敌人身上。"
    if "(qidx < GRID_CELLS)" in s:
        return "先检查格号是否在棋盘总范围内，防止数组访问越界。"
    if "(qidx != head_idx)" in s:
        return "禁止在玩家头上生成对象，否则玩家会无预警受伤。"
    if "!snake_map[qidx]" in s:
        return "禁止在玩家身体上生成对象，保证生成出来的位置公平且可见。"
    if "!food_occupies_index(qidx)" in s:
        return "禁止与已有食物重叠，避免一个格子同时显示多个含义。"
    if "!enemy_occupies_index(qidx)" in s:
        return "禁止与已有敌人重叠，避免敌人生成后立即互相碰撞。"
    if "enemy_spawn_valid = (qidx < GRID_CELLS) && (qx >= qlen - 1)" in s:
        return "敌人身体向左摆放，因此头部 x 坐标必须至少能容纳 qlen-1 个左侧身体段。"
    if "if ((sj < qlen) && ((qx < sj) || !spawn_cell_free(qidx - sj)))" in s:
        return "逐节检查敌人身体将占用的格子；任何一节越界或被占用，本次生成位置都无效。"
    if "enemy_pos[enemy_base + ej] <= enemy_scan_idx - ej" in s:
        return "敌人生成时头在 enemy_scan_idx，身体依次向左排开，初始方向设为向右就不会立刻回头撞身体。"
    if "enemy_pos[enemy_base + ej] <= INVALID_INDEX" in s:
        return "超过实际长度的数组槽写成无效格，避免渲染和碰撞把多余槽当成身体。"
    if "enemy_alive[spawn_slot] <= 1'b1" in s:
        return "找到合法位置后才激活敌人槽，渲染和碰撞逻辑才会开始处理它。"
    if "enemy_dir[spawn_slot] <= DIR_RIGHT" in s:
        return "敌人按向左排身体生成，所以初始向右移动能让身体自然跟在头后面。"
    if "enemy_len[spawn_slot] <= spawn_len" in s:
        return "把本次随机长度保存下来，后续渲染和移动只处理实际长度内的身体段。"

    if "food_valid <= 1'b1" in s:
        return "找到合法格后标记红食物存在，渲染和吃食物判断才会使用 food_idx。"
    if "yellow_food_valid <= 1'b1" in s:
        return "找到合法格后标记黄食物存在，避免同类食物再次生成。"
    if "green_food_valid <= 1'b1" in s:
        return "找到合法格后标记生命食物存在，玩家进入该格才会加生命。"
    if "food_valid <= 1'b0" in s:
        return "红食物被吃掉或游戏重置时清除存在标志，允许后续重新生成。"
    if "yellow_food_valid <= 1'b0" in s:
        return "黄食物被吃掉或游戏重置时清除存在标志，避免旧坐标继续生效。"
    if "green_food_valid <= 1'b0" in s:
        return "生命食物被吃掉或游戏重置时清除存在标志，避免重复加生命。"
    if "food_scan_idx <= food_scan_idx + 1'b1" in s:
        return "当前候选格被占用时继续向后扫描，直到找到空格或扫描完棋盘。"
    if "enemy_scan_idx <= enemy_scan_idx + 1'b1" in s:
        return "当前候选位置容不下敌人时继续扫描下一格，保证生成尽量成功。"

    if "nx = hx" in s:
        return "先把下一步坐标初始化为当前头坐标，再根据方向只修改 x 或 y。"
    if "ny = hy" in s:
        return "先保留当前 y 坐标，左右移动时 y 不变。"
    if "DIR_UP:    ny = hy - 1" in s:
        return "向上移动就是行号减 1；如果已经在第 0 行，后面边界判断会判定撞墙。"
    if "DIR_DOWN:  ny = hy + 1" in s:
        return "向下移动就是行号加 1；超过 ROWS-1 时会撞墙。"
    if "DIR_LEFT:  nx = hx - 1" in s:
        return "向左移动就是列号减 1；在最左侧继续左移会撞墙。"
    if "DIR_RIGHT: nx = hx + 1" in s:
        return "向右移动就是列号加 1；在最右侧继续右移会撞墙。"
    if "nidx = ny * COLS + nx" in s:
        return "把二维坐标转换为一维数组下标，公式是 行号*列数+列号，方便和 snake_map、食物、敌人数组比较。"
    if "if (slen == 0 || slen > MAXLEN)" in s:
        return "蛇长为 0 或超过数组容量说明状态异常，直接结束游戏避免继续访问非法数组。"
    if "else if (nx >= COLS || ny >= ROWS)" in s:
        return "下一步坐标超出棋盘宽高就是撞墙，需要扣生命或结束游戏。"
    if "if (invincible_ticks == 0)" in s:
        return "只有无敌时间结束后才允许再次受伤，防止一次碰撞连续扣多条命。"
    if "lives <= lives - 1'b1" in s:
        return "还有剩余生命时只扣一条命，让玩家可以继续游戏或重生。"
    if "lives <= 0" in s:
        return "生命耗尽时归零，状态显示和 SDK 监控都能看到游戏已经失败。"
    if "state <= 3" in s:
        return "把游戏状态切到 GAMEOVER；top.v 会据此清屏并显示 GAMEOVER 文字。"
    if "invincible_ticks <= INVINCIBLE_TICKS" in s:
        return "受伤后开启无敌倒计时，给玩家短暂保护时间。"
    if "invincible_blink <= 1'b0" in s:
        return "受伤瞬间先隐藏玩家一帧，配合后续翻转形成闪烁反馈。"
    if "invincible_blink <= ~invincible_blink" in s:
        return "无敌期间每个 tick 翻转显示状态，让玩家知道自己处于受伤保护期。"
    if "invincible_ticks <= invincible_ticks - 1'b1" in s:
        return "每个移动 tick 消耗一次无敌时间，直到恢复正常受伤判定。"

    # Display and VGA.
    if "H_VISIBLE = 640" in s:
        return "水平可见像素为 640，这是标准 VGA 640x480 模式的画面宽度。"
    if "V_VISIBLE = 480" in s:
        return "垂直可见像素为 480，这是标准 VGA 640x480 模式的画面高度。"
    if "H_TOTAL" in s and "=" in s:
        return "一行不仅包含可见像素，还包含前肩、同步和后肩，总长度用于控制显示器水平扫描节奏。"
    if "V_TOTAL" in s and "=" in s:
        return "一帧不仅包含可见行，还包含前肩、同步和后肩，总长度用于控制显示器垂直刷新节奏。"
    if "hcnt == H_TOTAL - 1" in s:
        return "水平计数到一行末尾后清零，并推动垂直计数进入下一行。"
    if "vcnt == V_TOTAL - 1" in s:
        return "垂直计数到一帧末尾后清零，下一次从屏幕左上角重新扫描。"
    if "gameover_box" in s and "pixel_x" in s:
        return "先限定 GAMEOVER 文字矩形区域，只有这个区域内才需要查询点阵字体。"
    if "gameover_text_x" in s and ">> 3" in s:
        return "文字被放大 8 倍显示，所以像素坐标右移 3 位还原成字体网格坐标。"
    if "gameover_pixel(gameover_text_x, gameover_text_y)" in s:
        return "根据放大后的字体坐标查询该点是否属于 GAMEOVER 字形。"

    # Buttons and seven-seg.
    if "btn_rise = btn_stable & ~btn_stable_d" in s:
        return "只在按键从未按到按下的瞬间产生事件，避免长按时连续触发方向或暂停。"
    if "btn_up_event" in s and "btn_rise" in s:
        return "把上升沿的 bit1 命名为上键事件，后面方向选择更直观。"
    if "dir_req = DIR_UP" in s:
        return "上键事件把请求方向设为向上，等下一个游戏 tick 生效。"
    if "dir_req_valid = 1'b1" in s:
        return "方向事件有效标志为 1，snake_game 才会接受本次方向请求。"
    if "btn_sync0 <= gpio_btn" in s:
        return "外部按键信号先打一拍进入像素时钟域，降低亚稳态风险。"
    if "btn_sync1 <= btn_sync0" in s:
        return "第二级同步进一步稳定按键信号，再交给去抖逻辑。"
    if "if (btn_sync1 == btn_stable)" in s:
        return "如果新采样值和稳定值相同，说明没有变化，去抖计数清零。"
    if "btn_debounce_cnt >= (BUTTON_DEBOUNCE_CLKS - 1)" in s:
        return "只有新电平持续足够久才更新稳定值，机械抖动不会误触发。"
    if "pause_reg <= ~pause_reg" in s:
        return "运行中按中心键翻转暂停状态，实现按一次暂停、再按一次继续。"
    if "reset_game_reg <= 1'b1" in s:
        return "GAMEOVER 时中心键产生复位脉冲，让 snake_game 清空状态等待重新开始。"
    if "score_capped" in s:
        return "数码管只显示 0-999，分数超过 999 时封顶显示，避免三位显示溢出。"
    if "score_hundreds" in s or "score_tens" in s or "score_ones" in s:
        return "把二进制分数拆成十进制位，才能逐位送到七段数码管显示。"
    if "active_enemy_limit = (elapsed_min >= 4'd2)" in s:
        return "游戏运行 2 分钟后把敌人上限从 3 提到 5，用时间推进提高难度。"
    if "sevenseg_scan <= sevenseg_scan + 1'b1" in s:
        return "快速轮流选择 8 个数码管位，人眼暂留会看到它们像同时点亮。"
    if "an = 8'b" in s:
        return "an 是低有效位选信号，只有对应的一位被拉低时该数码管点亮。"
    if "seg = sevenseg_pattern(sevenseg_digit)" in s:
        return "把当前要显示的十进制数字转换成七段管 a-g 的点亮模式。"

    # XDC constraints.
    if path_text and path_text.endswith(".xdc"):
        if s.startswith("set_property PACKAGE_PIN"):
            return "这行把逻辑端口接到板卡实际管脚；没有它，综合后的信号不会出现在 VGA 接口对应引脚上。"
        if s.startswith("set_property IOSTANDARD"):
            return "设置 3.3V LVCMOS 电平，和 Nexys DDR4 外设接口电气标准匹配。"
        if s.startswith("create_clock"):
            return "告诉 Vivado sys_clock 是 100MHz，后续时序分析才能判断电路是否跑得动。"
        if s.startswith("#"):
            return "约束文件说明，标出下面这些管脚属于哪一类外设信号。"

    # General but still meaningful fallbacks.
    if s.startswith("function "):
        return "把这段判断封装成函数，是为了在生成、移动、碰撞等多个位置复用同一套规则，避免逻辑不一致。"
    if s.startswith("assign "):
        return "这里使用组合逻辑输出，输入状态一变化，显示或控制信号就立即跟随变化。"
    if s.startswith("always @(posedge"):
        return "这里必须用时钟边沿更新，因为游戏状态、位置和计数器都要一拍一拍稳定推进。"
    if s.startswith("always @(*)"):
        return "这里是组合判断，只根据当前输入/状态决定输出，不需要额外记忆。"
    if s.startswith("case "):
        return "用 case 把编码值映射成对应行为，适合方向、状态、字符或颜色这种离散选择。"
    if s.startswith("for "):
        return "用循环遍历固定规模数组，硬件综合时会展开成并行/组合比较或多路更新逻辑。"
    if s in {"begin", "end", "endcase", "endfunction", "endmodule"}:
        return None
    if s.startswith("//"):
        return None

    # Avoid the old vague wording; leave low-value punctuation lines uncommented.
    return None


def annotate_existing_code(doc):
    path_text = ""
    title_text = ""
    title_re = re.compile(r"（([^（）]+):\d+-\d+）")
    code_re = re.compile(r"^\s*(\d+):\s?(.*)$")
    count = 0

    for para in doc.paragraphs:
        text = para.text
        title_match = title_re.search(text)
        if title_match:
            path_text = title_match.group(1)
            title_text = text
            continue

        code_match = code_re.match(text)
        if not code_match:
            continue

        line_no = int(code_match.group(1))
        code = code_match.group(2)
        explanation = explain_line(path_text, title_text, line_no, code)
        if not explanation:
            continue

        run = para.add_run(f"    {comment_prefix(path_text)}{explanation}")
        set_run_font(run, size=8.2, color=(0, 112, 192))
        count += 1

    return count


def add_annotated_code(doc, title, rel_path, start, end):
    path = ROOT / rel_path
    lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
    add_body(doc, f"{title}（{rel_path}:{start}-{end}，逻辑注释版）")
    for no in range(start, end + 1):
        if no < 1 or no > len(lines):
            continue
        code = lines[no - 1]
        p = doc.add_paragraph()
        set_para_spacing(p, after=0, line=1.0)
        shade_paragraph(p, "F2F2F2")
        r = p.add_run(f"{no:4d}: {code}")
        set_run_font(r, name="Consolas", size=8.3)
        explanation = explain_line(rel_path, title, no, code)
        if explanation:
            cr = p.add_run(f"    {comment_prefix(rel_path)}{explanation}")
            set_run_font(cr, size=8.1, color=(0, 112, 192))


def add_seed_section(doc):
    add_h1(doc, "12. seed 的输入来源与取值方式（逻辑补充）")
    add_body(
        doc,
        "这里的 seed 不是一个外部端口，也不是按键输入。它是几个随机相关函数的形参，作用是把同一套“随机输入”带入不同计算规则。"
        "调用这些函数时，实参通常来自 snake_game.v 内部的 16 位 LFSR：直接使用 lfsr，或把 lfsr 与当前位置、扫描位置、固定常数异或后再传入。"
    )
    add_body(
        doc,
        "为什么要这样做：LFSR 每个时钟都在变，但如果所有随机事件都直接使用同一个 lfsr，敌人长度、生成间隔、食物类型容易呈现固定关联。"
        "把 lfsr 与 enemy_scan_idx、food_scan_idx、head_idx 或常数异或，相当于为不同用途加入不同扰动，让“同一时刻的随机源”在不同功能上表现得更分散。"
    )
    add_table(
        doc,
        ["调用位置", "seed 实际取值", "为什么这样取", "得到什么效果"],
        [
            ["random_enemy_len(lfsr ^ enemy_scan_idx)", "当前 LFSR 与敌人生成扫描起点异或", "同一随机时刻下，不同生成位置会得到不同长度", "敌人长度在 4-8 之间变化，且不总和生成时刻强绑定。"],
            ["next_spawn_delay(lfsr ^ enemy_scan_idx)", "当前 LFSR 与敌人位置异或", "让下一次敌人出现时间受本次生成位置扰动", "敌人刷新节奏不会固定成机械周期。"],
            ["select_food_type(lfsr ^ head_idx)", "当前 LFSR 与玩家蛇头格号异或", "把玩家当前局面混入食物类型选择", "食物类型仍按 5:3:2，但序列不容易被固定复现。"],
            ["food_scan_idx <= lfsr % GRID_CELLS", "当前 LFSR 对 1200 取模", "把 16 位伪随机数压到棋盘格号范围内", "食物扫描从随机位置开始，避免总从左上角找空格。"],
            ["enemy_scan_idx <= lfsr % GRID_CELLS", "当前 LFSR 对 1200 取模", "把随机数转成候选生成格", "敌人从随机位置开始尝试生成。"],
            ["next_food_delay(lfsr ^ food_scan_idx)", "当前 LFSR 与食物实际位置异或", "让下一次食物刷新节奏受本次落点影响", "食物时间和位置都不固定。"],
            ["next_food_delay(16'hACE1) / next_spawn_delay(16'hACE1)", "复位固定非零常数", "保证刚复位时也有合法 seed，不依赖尚未滚动的随机序列", "开局倒计时能确定初始化。"],
            ["lfsr <= {lfsr[14:0], feedback}", "旧 lfsr 移位加抽头异或", "硬件中移位和异或资源少，适合 FPGA 生成伪随机数", "每拍产生新的 16 位 seed 来源。"],
        ],
    )

    add_h2(doc, "12.1 seed 相关代码逻辑注释版")
    add_annotated_code(doc, "食物类型 seed 如何变成 5:3:2 概率", "vga_snake.srcs/sources_1/new/snake_game.v", 250, 276)
    add_annotated_code(doc, "敌人长度和刷新间隔为什么这样使用 seed", "vga_snake.srcs/sources_1/new/snake_game.v", 512, 531)
    add_annotated_code(doc, "LFSR 为什么是 seed 的根来源", "vga_snake.srcs/sources_1/new/snake_game.v", 642, 645)
    add_annotated_code(doc, "LFSR 如何每拍产生新 seed", "vga_snake.srcs/sources_1/new/snake_game.v", 711, 712)
    add_annotated_code(doc, "敌人生成时 seed 如何结合扫描位置", "vga_snake.srcs/sources_1/new/snake_game.v", 976, 995)
    add_annotated_code(doc, "运行中食物和敌人如何从 lfsr 取随机起点", "vga_snake.srcs/sources_1/new/snake_game.v", 1220, 1239)


def main():
    doc = Document(str(SOURCE_DOC))
    added = annotate_existing_code(doc)
    add_seed_section(doc)
    OUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOC))
    print(OUT_DOC)
    print(f"semantic_comments={added}")


if __name__ == "__main__":
    main()
