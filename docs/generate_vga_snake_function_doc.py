from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "vga_snake_function_analysis.docx"


def set_run_font(run, name="Microsoft YaHei", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_para_spacing(paragraph, before=0, after=4, line=1.05):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=12, line=1.0)
    r = p.add_run(text)
    set_run_font(r, size=18, bold=True)


def add_h1(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=10, after=5, line=1.0)
    r = p.add_run(text)
    set_run_font(r, size=15, bold=True, color=(31, 78, 121))


def add_h2(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=7, after=4, line=1.0)
    r = p.add_run(text)
    set_run_font(r, size=12.5, bold=True, color=(47, 84, 150))


def add_body(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        set_para_spacing(p, after=2, line=1.05)
        r = p.add_run("• " + item)
        set_run_font(r, size=10.5)


def add_code(doc, title, rel_path, start, end):
    src_path = ROOT / rel_path
    lines = src_path.read_text(encoding="utf-8", errors="replace").splitlines()
    add_body(doc, f"{title}（{rel_path}:{start}-{end}）")
    for no in range(start, end + 1):
        if no < 1 or no > len(lines):
            continue
        p = doc.add_paragraph()
        set_para_spacing(p, after=0, line=1.0)
        shade_paragraph(p, "F2F2F2")
        r = p.add_run(f"{no:4d}: {lines[no - 1]}")
        set_run_font(r, name="Consolas", size=8.5)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(value))
            set_run_font(r, size=9.5)
    doc.add_paragraph()
    return table


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    add_title(doc, "VGA Snake 游戏具体功能实现说明与代码分析")
    add_body(doc, "本文档根据当前工程源码生成，重点说明敌人的随机长度、随机位置和随机运动，食物按比例随机出现，玩家身体变长与移动，GAMEOVER 显示，画面颜色控制，VGA 信号输出，碰撞判定，以及若干附加功能实现。")
    add_body(doc, f"工程目录：{ROOT}")

    add_h1(doc, "1. 工程总体结构")
    add_body(doc, "核心逻辑分为五层：板级封装 board_top.v、Vivado block design wrapper、游戏顶层 top.v、游戏状态机 snake_game.v、像素/颜色渲染 grid_renderer.v 与 VGA 时序 vga_timing.v。")
    add_table(
        doc,
        ["文件", "主要作用"],
        [
            ["vga_snake.srcs/sources_1/new/snake_game.v", "维护蛇、食物、敌人、分数、生命、状态机和碰撞；向渲染模块提供 cell_state。"],
            ["vga_snake.srcs/sources_1/new/grid_renderer.v", "把当前像素坐标转换为网格坐标，并按 cell_state 输出 RGB 颜色。"],
            ["vga_snake.srcs/sources_1/new/vga_timing.v", "产生 640x480@60Hz 的 hsync、vsync、active_video、pixel_x、pixel_y。"],
            ["vga_snake.srcs/sources_1/new/top.v", "连接按键、tick、snake_game、grid_renderer、vga_timing、GAMEOVER 覆盖显示、数码管。"],
            ["vga_snake.srcs/sources_1/new/board_top.v", "板级入口，连接 vga_snake_wrapper 与 VGA/数码管/按键/时钟/复位等外设。"],
            ["vga_snake.srcs/constrs_1/new/nexys_ddr4_vga.xdc", "把 hsync、vsync、vga_r/g/b、sys_clock、reset 等逻辑端口绑定到 Nexys DDR4 管脚。"],
        ],
    )
    add_code(doc, "snake_game 模块参数与状态编码", "vga_snake.srcs/sources_1/new/snake_game.v", 10, 49)
    add_body(doc, "cell_state 是整个显示链路的关键编码：snake_game 按每个格子返回 EMPTY、BODY、HEAD、不同食物、敌人头/身体；grid_renderer 再把这些状态翻译为 RGB。")

    add_h1(doc, "2. 随机数来源：LFSR 伪随机序列")
    add_body(doc, "硬件里通常不用软件 rand()。本工程在 snake_game.v 中使用 16 位 LFSR。每个时钟周期用反馈多项式推进一次，后续的敌人长度、敌人位置、敌人方向、食物种类和生成延迟都从 lfsr 的不同位或取模结果获得。")
    add_code(doc, "LFSR 初始化与推进", "vga_snake.srcs/sources_1/new/snake_game.v", 642, 645)
    add_code(doc, "LFSR 每周期更新", "vga_snake.srcs/sources_1/new/snake_game.v", 711, 712)
    add_body(doc, "分析：LFSR 是确定性伪随机序列。复位后种子固定为 16'hACE1，但 reset_game 时会与常数异或，游戏中又不断滚动，所以玩家体验上表现为随机。")

    add_h1(doc, "3. 敌人的随机长度、随机位置和随机运动")
    add_h2(doc, "3.1 随机长度")
    add_body(doc, "敌人长度由 ENEMY_MIN_LEN 和 ENEMY_MAX_LEN 限定，random_enemy_len(seed) 先计算长度跨度 span，再用 seed % span 落到范围内。当前 top.v 传入的范围是 4 到 8，因此敌人长度为 4、5、6、7、8 中的一个。")
    add_code(doc, "敌人长度与生成延迟函数", "vga_snake.srcs/sources_1/new/snake_game.v", 512, 525)
    add_body(doc, "例子：若 ENEMY_MIN_LEN=4、ENEMY_MAX_LEN=8，span=5；seed%5 为 0 时长度为 4，为 4 时长度为 8。")

    add_h2(doc, "3.2 随机位置")
    add_body(doc, "敌人生成不是直接把随机坐标写进去，而是先用 lfsr % GRID_CELLS 得到一个候选格 enemy_scan_idx，然后检查从这个格子向左摆放 qlen 个身体段是否全部合法。若不合法，enemy_scan_idx 逐格加 1 继续扫描，直到找到可用位置。")
    add_code(doc, "敌人生成位置合法性检查", "vga_snake.srcs/sources_1/new/snake_game.v", 379, 387)
    add_code(doc, "敌人按长度从候选格向左生成", "vga_snake.srcs/sources_1/new/snake_game.v", 976, 1000)
    add_body(doc, "分析：spawn_cell_free 会排除玩家头、玩家身体、食物和已有敌人；enemy_spawn_valid 还要求 qx >= qlen - 1，避免向左摆身体时跨出边界。因此“随机位置”不是盲放，而是随机起点加合法扫描。")

    add_h2(doc, "3.3 随机运动")
    add_body(doc, "每个敌人在移动 tick 中从 lfsr 的不同两位取 rand_dir。若随机方向正好与当前方向相反，则改为向右转，避免敌人直接 180 度回头。随后 enemy_move_valid 检查下一格是否在边界内、是否撞到敌人自己/其他敌人、是否撞到食物。")
    add_code(doc, "敌人移动合法性函数", "vga_snake.srcs/sources_1/new/snake_game.v", 342, 351)
    add_code(doc, "敌人随机方向、碰撞玩家与身体平移", "vga_snake.srcs/sources_1/new/snake_game.v", 1171, 1216)
    add_body(doc, "分析：enemy_pos[base] 是敌人头，enemy_pos[base+1..] 是身体。移动时从尾到头依次复制上一节位置，最后把 enemy_pos[base] 写成 enemy_next_idx，这和玩家蛇的“头前进、身体跟随”思路一致。")

    add_h1(doc, "4. 食物按比例随机出现")
    add_body(doc, "食物有三种：红色食物、黄色食物、蓝色生命食物。weighted_food_type(seed) 用 seed%10 划分概率桶：0-4 为红色，5-7 为黄色，8-9 为绿色/蓝色生命食物。因此理论比例是 5:3:2。")
    add_code(doc, "食物类型按 5:3:2 权重选择", "vga_snake.srcs/sources_1/new/snake_game.v", 250, 284)
    add_body(doc, "select_food_type 还做了兜底：如果按权重选中的那种食物已经存在，就优先改选当前未出现的类型。all_food_present 为真时不再生成，避免三种食物重复铺满。")
    add_code(doc, "食物生成请求、随机起点与空格扫描", "vga_snake.srcs/sources_1/new/snake_game.v", 714, 739)
    add_code(doc, "运行中按倒计时触发食物生成", "vga_snake.srcs/sources_1/new/snake_game.v", 1232, 1240)
    add_body(doc, "食物效果在玩家移动到 nidx 后判断：红色加 1 分并立即增长 1 节；黄色加 2 分并设置 growth_pending，使后续再额外长一节；蓝色生命食物不增长身体，而是在 lives < MAX_LIVES 时加生命。")
    add_code(doc, "吃到不同食物后的效果", "vga_snake.srcs/sources_1/new/snake_game.v", 1107, 1165)

    add_h1(doc, "5. 玩家身体变长和移动")
    add_body(doc, "玩家蛇的数据结构由两部分组成：snake_pos 环形数组记录从头到尾的格子序列，head_ptr/tail_ptr 指向头尾；snake_map 是位图，快速判断某个格子是否为身体。注意 snake_map 不包含当前头，只包含身体段。")
    add_code(doc, "蛇身存储结构", "vga_snake.srcs/sources_1/new/snake_game.v", 179, 197)
    add_code(doc, "开局初始化蛇头、两节身体和方向", "vga_snake.srcs/sources_1/new/snake_game.v", 740, 803)
    add_body(doc, "每个 tick 先根据 next_dir 算出下一格 nx/ny，再换算为 nidx。普通移动时：插入新头，旧头写入 snake_map，尾巴从 snake_map 清掉并移动 tail_ptr。吃红/黄食物时：插入新头但不删尾巴，所以 slen 增加。")
    add_code(doc, "根据方向计算下一格", "vga_snake.srcs/sources_1/new/snake_game.v", 1009, 1032)
    add_code(doc, "吃食物时立即增长", "vga_snake.srcs/sources_1/new/snake_game.v", 1118, 1138)
    add_code(doc, "普通移动或延迟增长", "vga_snake.srcs/sources_1/new/snake_game.v", 1143, 1159)
    add_body(doc, "关键点：普通移动允许蛇头走到当前尾巴的位置，因为尾巴在同一 tick 会离开，所以自撞判断写成 snake_map[nidx] && nidx != tailidx。")

    add_h1(doc, "6. 游戏结束后 GAMEOVER 的显示")
    add_body(doc, "游戏结束由 snake_game 把 state 置为 3。top.v 中为 GAMEOVER 定义了 5x7 点阵字体：gameover_letter_pixel 根据字符编号和字体内坐标返回该像素是否点亮，gameover_pixel 再把 G A M E O V E R 八个字符排列起来。")
    add_code(doc, "GAMEOVER 字体和字符排布函数", "vga_snake.srcs/sources_1/new/top.v", 110, 234)
    add_body(doc, "最终输出时，如果 state==3，top.v 不再使用 grid_renderer 的游戏画面颜色，而是把背景清黑，只在 GAMEOVER 点阵所在像素输出白色。")
    add_code(doc, "GAMEOVER 覆盖 VGA 输出", "vga_snake.srcs/sources_1/new/top.v", 508, 525)

    add_h1(doc, "7. 画面中颜色的控制")
    add_body(doc, "颜色控制分两步。第一步在 snake_game.v 中按优先级把当前 cell_x/cell_y 转换为 cell_state：玩家头、敌人头、食物、敌人身体、玩家身体、空格。第二步在 grid_renderer.v 中把 cell_state 映射成 3 位 R/G/B。")
    add_code(doc, "格子状态优先级输出", "vga_snake.srcs/sources_1/new/snake_game.v", 571, 600)
    add_code(doc, "cell_state 到 RGB 的颜色表", "vga_snake.srcs/sources_1/new/grid_renderer.v", 42, 91)
    add_table(
        doc,
        ["cell_state", "含义", "RGB", "显示颜色"],
        [
            ["000", "空格", "000/000/000，边框 010/010/010", "黑色背景，灰色网格线"],
            ["001", "玩家身体", "000/111/000", "绿色"],
            ["010", "玩家头", "111/111/111", "白色"],
            ["011", "红色食物", "111/000/000", "红色"],
            ["100", "黄色食物", "111/111/000", "黄色"],
            ["101", "生命食物", "000/000/111", "蓝色"],
            ["110", "敌人头", "111/000/111", "品红"],
            ["111", "敌人身体", "011/000/011", "暗品红"],
        ],
    )
    add_body(doc, "分析：grid_renderer 在 active_video 为 0 时输出黑色，避免消隐区出现无效颜色；GAMEOVER 状态又在 top.v 末尾覆盖 RGB，因此 GAMEOVER 的显示优先级最高。")

    add_h1(doc, "8. 信号如何传输到显示屏")
    add_body(doc, "显示链路为：sys_clock 进入 Vivado block design 的 Clocking Wizard，产生 clk_pix；top.v 中 vga_timing 用 clk_pix 产生 hsync/vsync/active/pixel_x/pixel_y；grid_renderer 根据像素坐标查询 snake_game 的 cell_state 并输出 rgb_r/g/b；top.v 最终把 rgb 或 GAMEOVER 文字输出到 vga_r/g/b；board_top 和 XDC 再把这些逻辑端口连到实际 VGA 管脚。")
    add_code(doc, "VGA 640x480 时序参数与同步信号", "vga_snake.srcs/sources_1/new/vga_timing.v", 34, 94)
    add_code(doc, "top 中连接时序、游戏和渲染模块", "vga_snake.srcs/sources_1/new/top.v", 310, 342)
    add_code(doc, "snake_game 与 grid_renderer 的端口连接", "vga_snake.srcs/sources_1/new/top.v", 457, 506)
    add_code(doc, "板级封装把 VGA 信号接到 wrapper", "vga_snake.srcs/sources_1/new/board_top.v", 3, 35)
    add_code(doc, "VGA 逻辑端口到 Nexys DDR4 管脚", "vga_snake.srcs/constrs_1/new/nexys_ddr4_vga.xdc", 4, 41)
    add_body(doc, "补充：综合后的 block design 中 top_0 的 clk_pix 连接到 clk_wiz_1_clk_out2，top_0_hsync/vga_r/vga_g/vga_b/vsync 再直接赋给外部端口，因此显示信号从游戏逻辑一路传到板卡 VGA 端口。")

    add_h1(doc, "9. 碰撞判定")
    add_body(doc, "碰撞判定集中在玩家 tick 更新和敌人 tick 更新中。玩家先算下一格 nidx，再按顺序判断边界、自身、敌人头、敌人身体、食物；敌人移动时再判断是否撞到玩家。")
    add_code(doc, "玩家撞墙、自撞与生命处理", "vga_snake.srcs/sources_1/new/snake_game.v", 1034, 1074)
    add_code(doc, "玩家撞敌人头、敌人身体和食物", "vga_snake.srcs/sources_1/new/snake_game.v", 1075, 1116)
    add_code(doc, "敌人移动撞玩家", "vga_snake.srcs/sources_1/new/snake_game.v", 1189, 1208)
    add_table(
        doc,
        ["碰撞类型", "判定条件", "结果"],
        [
            ["撞墙", "nx>=COLS 或 ny>=ROWS", "无敌期外扣生命；生命为 1 时 state<=3。"],
            ["自撞", "nidx==head_idx 或 snake_map[nidx] 且 nidx!=tailidx", "扣生命或 GAMEOVER；若还有生命则进入重生流程。"],
            ["撞敌人头", "enemy_head_at_index(nidx)", "玩家受伤，生命不足时 GAMEOVER。"],
            ["撞敌人身体", "enemy_pos[enemy_base+ej]==nidx，ej>=1", "击杀该敌人，后续加 5 分。"],
            ["吃食物", "nidx 等于 food_idx/yellow_food_idx/green_food_idx", "红/黄增长和加分；蓝色增加生命。"],
            ["敌人撞玩家", "enemy_next_idx==nidx/head_idx 或撞 snake_map", "玩家受伤或 GAMEOVER。"],
            ["生成冲突", "spawn_cell_free、food_occupies_index、enemy_occupies_index", "不在被占格生成，继续扫描下一个格子。"],
        ],
    )
    add_body(doc, "分析：碰撞检测把“玩家本 tick 的新头 nidx”和“旧头 head_idx”“身体位图 snake_map”“尾巴 tailidx”同时考虑，解决了蛇类游戏里常见的尾巴同 tick 移走问题。")

    add_h1(doc, "10. 其他具体功能例子")
    add_h2(doc, "10.1 按键同步、去抖、方向优先级、暂停/复位")
    add_body(doc, "top.v 对五个按键先做两级同步，再用 BUTTON_DEBOUNCE_CLKS 计数去抖。方向键只在上升沿产生一次 dir_req_valid；中键在运行时切换 pause，在 GAMEOVER 时产生 reset_game。")
    add_code(doc, "按键事件、去抖和暂停复位", "vga_snake.srcs/sources_1/new/top.v", 79, 89)
    add_code(doc, "方向优先级与去抖", "vga_snake.srcs/sources_1/new/top.v", 236, 304)

    add_h2(doc, "10.2 移动 tick 控制速度")
    add_body(doc, "像素时钟是 25 MHz，但蛇不可能每个像素周期都移动。top.v 用 tick_cnt 分频，TICK_DIV=2500000 时约 10 Hz，即每秒移动约 10 格。")
    add_code(doc, "移动 tick 产生器", "vga_snake.srcs/sources_1/new/top.v", 320, 342)

    add_h2(doc, "10.3 分数、生命、时间、敌人上限和数码管")
    add_body(doc, "score/lives/state 由 snake_game 输出；top.v 把生命、时间、分数拆成十进制位，通过 sevenseg_scan 轮流点亮 8 位数码管。同时 elapsed_min>=2 时 active_enemy_limit 从 3 增到 5，形成随时间提高难度的功能。")
    add_code(doc, "时间、分数拆位、敌人上限和数码管扫描", "vga_snake.srcs/sources_1/new/top.v", 347, 447)
    add_code(doc, "状态打包给 GPIO/MicroBlaze", "vga_snake.srcs/sources_1/new/top.v", 484, 484)

    add_h2(doc, "10.4 无敌闪烁与重生")
    add_body(doc, "玩家受伤但生命未耗尽时会设置 invincible_ticks，并让 invincible_blink 翻转；cell_state 输出阶段用 show_player 控制是否显示玩家，实现闪烁。自撞后还会 pending_player_respawn，扫描安全位置重建蛇身。")
    add_code(doc, "无敌计数与闪烁", "vga_snake.srcs/sources_1/new/snake_game.v", 1018, 1023)
    add_code(doc, "受伤后进入重生扫描", "vga_snake.srcs/sources_1/new/snake_game.v", 1060, 1073)

    add_h1(doc, "11. 总结")
    add_body(doc, "本工程的实现思路可以概括为：snake_game 维护离散网格世界，grid_renderer 把网格世界转换为像素颜色，vga_timing 提供显示器需要的扫描时序，top/board_top 完成按键、数码管、GAMEOVER 覆盖和板级连接。随机功能都来自 LFSR；碰撞功能都围绕下一格 nidx、身体位图、敌人数组和食物坐标展开；颜色和显示则通过 cell_state 到 RGB 的两级映射完成。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
