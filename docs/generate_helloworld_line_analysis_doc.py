from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "vga_snake.sdk" / "vga_snake_app" / "src" / "helloworld.c"
OUT = ROOT / "docs" / "helloworld_c_line_by_line_analysis.docx"


def set_run_font(run, name="Microsoft YaHei", size=9.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, font="Microsoft YaHei", size=9.0, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, name=font, size=size, bold=bold, color=color)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    set_run_font(r, size=17, bold=True, color=(31, 78, 121))


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=13, bold=True, color=(47, 84, 150))


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    set_run_font(r, size=10.5)


def add_small_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "D9EAF7")
        set_cell_text(cell, header, size=9.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), size=9.0)
    doc.add_paragraph()
    return table


EXPLANATIONS = {
    1: "文件头注释开始，使用一整块星号注释说明本文件用途。",
    2: "注释中的空行，用来让文件头更清晰。",
    3: "说明程序名称：这是 VGA Snake 的 SDK 监控应用。",
    4: "注释中的空行，分隔标题和详细说明。",
    5: "说明蛇游戏逻辑和 VGA 渲染都在 RTL 硬件逻辑中完成，不在 C 程序中实现游戏本体。",
    6: "说明 MicroBlaze 通过 AXI GPIO 读取板上按键和游戏状态。",
    7: "说明程序通过 UART 打印信息，便于在 SDK 串口终端验证软硬件连接。",
    8: "注释中的空行。",
    9: "文件头注释结束。",
    10: "空行，分隔文件注释和头文件包含。",
    11: "包含 platform.h，获得 init_platform() 和 cleanup_platform() 的声明，用于初始化/清理 BSP 平台。",
    12: "包含 sleep.h，获得 usleep()，后面循环中用它控制串口打印采样周期。",
    13: "包含 xgpio.h，获得 XGpio 类型以及 AXI GPIO 初始化、读写和方向配置函数。",
    14: "包含 xil_printf.h，使用 xil_printf() 通过 UART 输出调试文本。",
    15: "包含 xparameters.h，使用 Vivado/SDK 自动生成的硬件参数，例如 AXI GPIO 设备 ID 和地址。",
    16: "包含 xstatus.h，使用 XST_SUCCESS、XST_FAILURE 等 Xilinx 状态码。",
    17: "空行，分隔 include 和宏定义。",
    18: "把硬件参数 XPAR_AXI_GPIO_0_DEVICE_ID 封装为 GPIO_DEVICE_ID，后面初始化 AXI GPIO 时使用。",
    19: "空行，分隔设备 ID 和通道编号。",
    20: "定义 AXI GPIO 通道 1 为按键输入通道。",
    21: "定义 AXI GPIO 通道 2 为游戏状态输入通道。",
    22: "空行，分隔通道宏和按键位掩码。",
    23: "定义中心按键位掩码，bit0 对应 CENTER。",
    24: "定义上方向按键位掩码，bit1 对应 UP。",
    25: "定义左方向按键位掩码，bit2 对应 LEFT。",
    26: "定义右方向按键位掩码，bit3 对应 RIGHT。",
    27: "定义下方向按键位掩码，bit4 对应 DOWN。",
    28: "定义总按键掩码 0x1F，即只保留低 5 位按键信号。",
    29: "空行，分隔按键位定义和状态位定义。",
    30: "定义分数掩码，低 16 位 status[15:0] 表示 score。",
    31: "定义状态字段右移位数，游戏状态位于 status[17:16]。",
    32: "定义状态字段掩码 0x00030000，用于提取 status[17:16]。",
    33: "定义生命数字段右移位数，生命位于 status[21:18]。",
    34: "定义生命数字段掩码 0x003C0000，用于提取 status[21:18]。",
    35: "空行，分隔状态位域和状态枚举值。",
    36: "定义游戏空闲状态 IDLE 的编码为 0。",
    37: "定义游戏运行状态 RUNNING 的编码为 1。",
    38: "定义游戏暂停状态 PAUSE 的编码为 2。",
    39: "定义游戏结束状态 GAMEOVER 的编码为 3。",
    40: "空行，分隔宏定义和全局变量。",
    41: "定义静态全局 XGpio 实例 Gpio，保存 AXI GPIO 驱动实例状态；static 限定它只在本文件可见。",
    42: "空行，分隔全局变量和函数。",
    43: "定义 state_name 函数，把数值状态转换为字符串，便于串口打印。",
    44: "state_name 函数体开始。",
    45: "根据传入的 state 使用 switch 分支选择对应字符串。",
    46: "当 state 等于空闲状态编码时进入该分支。",
    47: "返回字符串 IDLE。",
    48: "当 state 等于运行状态编码时进入该分支。",
    49: "返回字符串 RUNNING。",
    50: "当 state 等于暂停状态编码时进入该分支。",
    51: "返回字符串 PAUSE。",
    52: "当 state 等于游戏结束状态编码时进入该分支。",
    53: "返回字符串 GAMEOVER。",
    54: "默认分支，用来处理未知状态编码。",
    55: "返回 UNKNOWN，避免串口打印空指针或无意义内容。",
    56: "switch 语句结束。",
    57: "state_name 函数结束。",
    58: "空行，分隔函数。",
    59: "定义 status_state 函数，从完整 status 字中提取游戏状态字段。",
    60: "status_state 函数体开始。",
    61: "先用 STATUS_STATE_MASK 保留 bit17:16，再右移 16 位得到 0、1、2、3 的状态值。",
    62: "status_state 函数结束。",
    63: "空行，分隔函数。",
    64: "定义 is_quiet_state 函数，用来判断哪些状态不需要连续刷屏打印。",
    65: "is_quiet_state 函数体开始。",
    66: "判断 state 是否为 IDLE；空闲状态变化少，属于静默状态。",
    67: "继续判断 state 是否为 PAUSE；暂停状态也属于静默状态。",
    68: "继续判断 state 是否为 GAMEOVER；游戏结束状态也属于静默状态。",
    69: "is_quiet_state 函数结束。",
    70: "空行，分隔函数。",
    71: "定义 print_buttons 函数，把按键位图翻译成可读的按键名称。",
    72: "print_buttons 函数体开始。",
    73: "先输出固定前缀 buttons:。",
    74: "空行，分隔前缀输出和具体按键判断。",
    75: "判断 UP 位是否为 1。",
    76: "如果按下 UP，则在串口输出 UP。",
    77: "UP 判断结束。",
    78: "判断 DOWN 位是否为 1。",
    79: "如果按下 DOWN，则输出 DOWN。",
    80: "DOWN 判断结束。",
    81: "判断 LEFT 位是否为 1。",
    82: "如果按下 LEFT，则输出 LEFT。",
    83: "LEFT 判断结束。",
    84: "判断 RIGHT 位是否为 1。",
    85: "如果按下 RIGHT，则输出 RIGHT。",
    86: "RIGHT 判断结束。",
    87: "判断 CENTER 位是否为 1。",
    88: "如果按下 CENTER，则输出 CENTER。",
    89: "CENTER 判断结束。",
    90: "如果低 5 位按键全部为 0，说明没有按键按下。",
    91: "输出 none，表示当前无按键输入。",
    92: "无按键判断结束。",
    93: "print_buttons 函数结束。",
    94: "空行，分隔函数。",
    95: "定义 print_status 函数，统一打印游戏状态、分数、生命和按键状态。",
    96: "print_status 函数体开始。",
    97: "从 status 低 16 位提取 score。",
    98: "从 status 的 bit17:16 提取 state。",
    99: "从 status 的 bit21:18 提取 lives。",
    100: "空行，分隔数据解析和串口输出。",
    101: "用 xil_printf 输出状态名、分数和生命数，末尾保留空格让按键信息接在同一行。",
    102: "传入 state_name(state)、score、lives，分别填充上一行格式字符串中的 %s 和 %u。",
    103: "调用 print_buttons(buttons)，继续输出当前按键名称。",
    104: "输出回车换行，结束本次状态打印。",
    105: "print_status 函数结束。",
    106: "空行，分隔辅助函数和 main 函数。",
    107: "定义 main 函数，是裸机 MicroBlaze 程序的入口。",
    108: "main 函数体开始。",
    109: "定义 status 变量，用来保存函数返回状态码。",
    110: "定义 buttons 变量，保存当前采样到的按键低 5 位。",
    111: "定义 game_status 变量，保存从 AXI GPIO 状态通道读到的完整游戏状态字。",
    112: "定义 sampled_buttons 保存上一轮采样的按键，用于检测新按下的边沿。",
    113: "定义 last_buttons 并初始化为全 1，保证第一次循环认为按键发生变化，从而打印一次。",
    114: "定义 last_status 并初始化为全 1，保证第一次循环认为状态发生变化，从而打印一次。",
    115: "定义 last_quiet_state，记录上一次已经打印过的静默状态。",
    116: "定义 quiet_state_printed 标志，表示当前静默状态是否已经打印过。",
    117: "空行，分隔变量定义和初始化调用。",
    118: "初始化 BSP 平台：打开缓存、初始化 UART 等，使后续 xil_printf 可用。",
    119: "空行。",
    120: "空行，代码排版留白，对程序逻辑没有影响。",
    121: "空行，代码排版留白，对程序逻辑没有影响。",
    122: "调用 XGpio_Initialize，根据 GPIO_DEVICE_ID 初始化 AXI GPIO 驱动实例 Gpio。",
    123: "判断 AXI GPIO 初始化是否失败。",
    124: "若初始化失败，通过 UART 打印错误信息和返回状态码。",
    125: "初始化失败时执行平台清理，主要是关闭缓存等。",
    126: "返回 XST_FAILURE，通知启动环境程序执行失败。",
    127: "初始化失败分支结束。",
    128: "空行，分隔初始化检查和通道方向配置。",
    129: "设置 AXI GPIO 通道 1 的低 5 位为输入方向，用于读取五个按键；XGpio 中方向位为 1 表示输入。",
    130: "设置 AXI GPIO 通道 2 的 score/state/lives 位为输入方向，用于读取 RTL 打包出来的游戏状态。",
    131: "空行，分隔硬件配置和提示输出。",
    132: "串口打印启动提示，说明 AXI GPIO 已准备好并开始监控按键与游戏状态。",
    133: "空行，分隔初始化阶段和主循环。",
    134: "进入 while(1) 无限循环，裸机监控程序会一直运行，不主动退出。",
    135: "定义本轮解析出的 state 临时变量。",
    136: "定义 button_press，用于保存新按下的按键边沿。",
    137: "定义 should_print，控制本轮是否需要串口打印。",
    138: "空行，分隔循环局部变量和采样操作。",
    139: "从 AXI GPIO 通道 1 读取按键，并用 BTN_MASK 只保留低 5 位有效按键。",
    140: "从 AXI GPIO 通道 2 读取游戏状态字，包括 score、state、lives。",
    141: "调用 status_state 从 game_status 中提取当前游戏状态。",
    142: "检测新按下的按键：当前为 1 且上一轮 sampled_buttons 为 0 的位会保留下来。",
    143: "空行，分隔采样和打印条件判断。",
    144: "默认打印条件：只要按键或状态相对上一次打印值发生变化，就准备打印。",
    145: "空行，分隔普通打印条件和静默状态特殊处理。",
    146: "判断当前是否为静默状态 IDLE/PAUSE/GAMEOVER。",
    147: "如果进入了新的静默状态，与上一次静默状态不同，则需要允许重新打印一次。",
    148: "清除 quiet_state_printed，让新的静默状态可以打印首条信息。",
    149: "新的静默状态判断结束。",
    150: "静默状态下只在尚未打印过，或检测到有新按键按下时打印，避免串口反复刷同一状态。",
    151: "如果不是静默状态，即 RUNNING 状态。",
    152: "清空 last_quiet_state，表示当前不处在静默状态。",
    153: "清除 quiet_state_printed，方便下次进入静默状态时重新打印。",
    154: "静默状态分支结束。",
    155: "空行，分隔打印条件计算和实际打印。",
    156: "如果本轮 should_print 为真，则执行串口输出。",
    157: "调用 print_status 打印当前状态、分数、生命和按键。",
    158: "打印后，如果当前仍是静默状态，需要更新静默状态记录。",
    159: "记录本次已经打印过的静默状态编号。",
    160: "置 quiet_state_printed 为 1，防止同一个静默状态下反复打印。",
    161: "静默状态记录更新结束。",
    162: "更新 last_buttons，作为下次判断按键是否变化的基准。",
    163: "更新 last_status，作为下次判断游戏状态是否变化的基准。",
    164: "should_print 分支结束。",
    165: "空行，分隔打印逻辑和循环尾部处理。",
    166: "把当前按键保存到 sampled_buttons，用于下一轮检测 button_press 上升沿。",
    167: "延时 100000 微秒，即 0.1 秒；控制采样和打印频率，避免串口输出过快。",
    168: "while 循环体结束，随后回到第 134 行继续监控。",
    169: "空行；由于 while(1) 不退出，下面两行正常运行时不会执行。",
    170: "调用 cleanup_platform 清理平台资源；理论上用于退出前关闭缓存，但在无限循环程序中通常不可达。",
    171: "返回 XST_SUCCESS；同样通常不可达，只是保持 main 函数完整规范。",
    172: "main 函数结束。",
}


def build_doc():
    lines = SRC.read_text(encoding="utf-8", errors="replace").splitlines()

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Pt(42)
    section.bottom_margin = Pt(42)
    section.left_margin = Pt(42)
    section.right_margin = Pt(42)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    add_title(doc, "SDK helloworld.c 逐行代码分析")
    add_body(doc, f"源码文件：{SRC}")
    add_body(doc, "该 C 程序运行在 MicroBlaze 裸机环境中，主要作用不是实现贪吃蛇游戏本体，而是通过 AXI GPIO 读取硬件 RTL 中的按键和游戏状态，并通过 UART 串口输出，方便在 SDK 中验证软硬件系统是否连通。")

    add_h1(doc, "1. 程序总体功能")
    add_small_table(
        doc,
        ["功能点", "说明"],
        [
            ["AXI GPIO 初始化", "使用 XGpio_Initialize(&Gpio, GPIO_DEVICE_ID) 初始化硬件 GPIO 外设。"],
            ["按键读取", "通道 1 读取 5 个按键，低 5 位分别对应 CENTER/UP/LEFT/RIGHT/DOWN。"],
            ["状态读取", "通道 2 读取 RTL 打包的状态字：score、state、lives。"],
            ["状态解析", "通过掩码和移位从 32 位 status 中取出不同字段。"],
            ["串口输出", "使用 xil_printf 打印状态、分数、生命和按键。"],
            ["降低刷屏", "IDLE/PAUSE/GAMEOVER 属于静默状态，只在进入状态或有新按键时打印。"],
        ],
    )

    add_h1(doc, "2. 硬件参数与位域")
    add_small_table(
        doc,
        ["项目", "当前工程中的值/含义"],
        [
            ["XPAR_AXI_GPIO_0_DEVICE_ID", "0，来自 xparameters.h，用于定位 AXI GPIO 驱动配置。"],
            ["XPAR_AXI_GPIO_0_BASEADDR", "0x40000000，AXI GPIO 的基地址。"],
            ["XPAR_AXI_GPIO_0_IS_DUAL", "1，表示 AXI GPIO 有两个通道。"],
            ["XPAR_AXI_UARTLITE_0_BASEADDR", "0x40600000，UARTLite 基地址，也是 STDOUT_BASEADDRESS。"],
            ["GPIO_CH_BUTTONS", "1，按键输入通道。"],
            ["GPIO_CH_STATUS", "2，游戏状态输入通道。"],
            ["status[15:0]", "score，游戏分数。"],
            ["status[17:16]", "state，0=IDLE，1=RUNNING，2=PAUSE，3=GAMEOVER。"],
            ["status[21:18]", "lives，生命数。"],
        ],
    )

    add_h1(doc, "3. 逐行代码分析")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["行号", "代码", "逐行解释"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "D9EAF7")
        set_cell_text(cell, header, size=9.5, bold=True)

    for no, line in enumerate(lines, 1):
        cells = table.add_row().cells
        set_cell_text(cells[0], str(no), size=8.5)
        code_text = line if line else "（空行）"
        set_cell_text(cells[1], code_text, font="Consolas", size=8.2)
        explanation = EXPLANATIONS.get(no, "该行用于维持 C 语言语法结构或代码排版。")
        set_cell_text(cells[2], explanation, size=8.8)

    add_h1(doc, "4. 关键流程总结")
    add_body(doc, "程序启动后先初始化平台和 AXI GPIO，然后把两个 GPIO 通道都配置为输入。主循环每 0.1 秒读取一次按钮和状态字，解析出 state/score/lives，再根据是否变化决定是否通过 UART 打印。")
    add_body(doc, "它与 RTL 游戏逻辑的关系是“监控和验证”：top.v/board design 负责把游戏状态打包到 AXI GPIO，helloworld.c 只负责读取并显示这些状态，因此 C 程序不会直接控制 VGA 画面和蛇的移动逻辑。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
